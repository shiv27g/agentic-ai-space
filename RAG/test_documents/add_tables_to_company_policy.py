"""
Add dummy tables to 11_company_policy.docx for the DOCX load demo.
Run once from Week-4-RAG: python test_documents/add_tables_to_company_policy.py
"""

import os
from pathlib import Path

try:
    import docx
    from docx import Document
except ImportError:
    print("pip install python-docx")
    raise

# Path to the company policy docx (same folder as this script)
SCRIPT_DIR = Path(__file__).resolve().parent
DOCX_PATH = SCRIPT_DIR / "11_company_policy.docx"
BACKUP_PATH = SCRIPT_DIR / "11_company_policy_backup.docx"


def add_dummy_tables():
    if not DOCX_PATH.exists():
        print(f"Not found: {DOCX_PATH}")
        return

    doc = Document(str(DOCX_PATH))

    # Add a heading before the tables
    doc.add_paragraph()
    doc.add_heading("Policy Summaries (Reference Tables)", level=2)

    # Table 1: Leave policy
    doc.add_paragraph()
    doc.add_paragraph("Leave entitlement by type:")
    t1 = doc.add_table(rows=4, cols=3)
    t1.style = "Table Grid"
    rows = [
        ["Leave Type", "Days/Year", "Notes"],
        ["Annual", "20", "Accrues monthly"],
        ["Sick", "10", "Medical certificate after 3 days"],
        ["Personal", "5", "No carry-over"],
    ]
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            t1.rows[i].cells[j].text = cell_text

    # Table 2: Expense limits
    doc.add_paragraph()
    doc.add_paragraph("Expense approval limits (USD):")
    t2 = doc.add_table(rows=4, cols=3)
    t2.style = "Table Grid"
    rows = [
        ["Limit", "Approver", "Requires receipt"],
        ["0 - 500", "Manager", "Yes"],
        ["500 - 5000", "Director", "Yes"],
        ["5000+", "VP / CFO", "Yes + quote"],
    ]
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            t2.rows[i].cells[j].text = cell_text

    # Table 3: Approval matrix
    doc.add_paragraph()
    doc.add_paragraph("Document approval matrix:")
    t3 = doc.add_table(rows=4, cols=3)
    t3.style = "Table Grid"
    rows = [
        ["Document type", "Approver", "SLA (days)"],
        ["Policy change", "Legal + HR", "5"],
        ["Contract", "Legal", "3"],
        ["Budget", "Finance", "2"],
    ]
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            t3.rows[i].cells[j].text = cell_text

    # Optional backup
    if not BACKUP_PATH.exists():
        import shutil
        shutil.copy(DOCX_PATH, BACKUP_PATH)
        print(f"Backup saved: {BACKUP_PATH}")

    doc.save(str(DOCX_PATH))
    print(f"Updated: {DOCX_PATH}")
    print("Added 3 dummy tables (Leave, Expense limits, Approval matrix).")
    print("Run: python load_docx_demo.py (from Week-4-RAG) to demo extraction.")


if __name__ == "__main__":
    add_dummy_tables()
