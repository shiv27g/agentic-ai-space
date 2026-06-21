"""
Utility script to load test documents with metadata for use in demos.

Supports native formats:
- .txt, .py: plain text (read directly)
- .docx: Word (python-docx)
- .html: HTML (BeautifulSoup or strip tags)
"""

import json
from pathlib import Path
from langchain_core.documents import Document
from typing import List, Dict, Optional

TEST_DOCUMENTS_DIR = Path(__file__).parent
METADATA_FILE = TEST_DOCUMENTS_DIR / "document_metadata.json"


def _read_content_by_format(file_path: Path, filename: str) -> str:
    """Read document content using the appropriate loader for file format."""
    suffix = file_path.suffix.lower()
    if suffix in (".txt", ".py", ".md"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    if suffix == ".docx":
        try:
            import docx
            doc = docx.Document(str(file_path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise FileNotFoundError(
                f"Cannot load {filename}: python-docx not installed. pip install python-docx"
            )
    if suffix in (".html", ".htm"):
        with open(file_path, "r", encoding="utf-8") as f:
            html = f.read()
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            import re
            return re.sub(r"<[^>]+>", " ", html).strip()
    # Fallback: try text
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def load_document(filename: str) -> Document:
    """
    Load a single document with its metadata.
    Uses format-appropriate loader: .txt/.py read as text, .docx via python-docx, .html via BeautifulSoup.

    Args:
        filename: Name of the document file (e.g., '01_financial_report.txt', '11_company_policy.docx')

    Returns:
        LangChain Document with content and metadata
    """
    file_path = TEST_DOCUMENTS_DIR / filename
    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {filename}")

    content = _read_content_by_format(file_path, filename)

    with open(METADATA_FILE, "r", encoding="utf-8") as f:
        all_metadata = json.load(f)
    metadata = all_metadata.get(filename, {"source_file": filename})

    return Document(page_content=content, metadata=metadata)


def load_all_documents() -> List[Document]:
    """
    Load all test documents with their metadata.
    
    Returns:
        List of LangChain Documents
    """
    documents = []
    
    # Load metadata to get list of files
    with open(METADATA_FILE, 'r', encoding='utf-8') as f:
        all_metadata = json.load(f)
    
    # Load each document
    for filename in all_metadata.keys():
        try:
            doc = load_document(filename)
            documents.append(doc)
        except FileNotFoundError:
            print(f"Warning: File not found: {filename}")
            continue
    
    return documents


def load_documents_by_filter(**filters) -> List[Document]:
    """
    Load documents matching specific metadata filters.
    
    Args:
        **filters: Metadata filters (e.g., department='finance', access_level='confidential')
    
    Returns:
        List of matching LangChain Documents
    """
    all_docs = load_all_documents()
    filtered_docs = []
    
    for doc in all_docs:
        match = True
        for key, value in filters.items():
            if doc.metadata.get(key) != value:
                match = False
                break
        if match:
            filtered_docs.append(doc)
    
    return filtered_docs


def load_documents_by_department(department: str) -> List[Document]:
    """Load all documents for a specific department."""
    return load_documents_by_filter(department=department)


def load_documents_by_access_level(access_level: str) -> List[Document]:
    """
    Load documents accessible at a given access level.
    Access levels are hierarchical: public < internal < confidential < secret
    
    Args:
        access_level: The user's access level (what they can access)
    
    Returns:
        Documents that the user with this access level can see
    """
    hierarchy = {
        'public': ['public'],
        'internal': ['public', 'internal'],
        'confidential': ['public', 'internal', 'confidential'],
        'secret': ['public', 'internal', 'confidential', 'secret']
    }
    
    allowed_levels = hierarchy.get(access_level, ['public'])
    all_docs = load_all_documents()
    
    # Return documents whose access_level is in the allowed list
    return [doc for doc in all_docs if doc.metadata.get('access_level') in allowed_levels]


def get_test_queries() -> Dict[str, str]:
    """
    Get a dictionary of test queries for different scenarios.
    
    Returns:
        Dictionary mapping scenario names to queries
    """
    return {
        'exact_term_match': 'AWS EC2 pricing November 2024',
        'semantic_match': 'cloud cost optimization strategies',
        'hybrid_test': 'What was Q4 revenue?',
        'access_control_test': 'What are the salary ranges?',
        'chunking_test': 'What is the damage cap in Section 7.3?',
        'embeddings_test': 'Where did the cat sit?',
    }


if __name__ == "__main__":
    # Example usage
    print("Loading all documents...")
    all_docs = load_all_documents()
    print(f"Loaded {len(all_docs)} documents\n")
    
    print("Sample document:")
    doc = all_docs[0]
    print(f"  File: {doc.metadata.get('source_file')}")
    print(f"  Department: {doc.metadata.get('department')}")
    print(f"  Access Level: {doc.metadata.get('access_level')}")
    print(f"  Content preview: {doc.page_content[:100]}...\n")
    
    print("Filtering by department='finance':")
    finance_docs = load_documents_by_department('finance')
    print(f"  Found {len(finance_docs)} documents")
    for doc in finance_docs:
        print(f"    - {doc.metadata.get('source_file')}")
    
    print("\nFiltering by access_level='confidential' (user can see public, internal, confidential):")
    confidential_docs = load_documents_by_access_level('confidential')
    print(f"  Found {len(confidential_docs)} documents")
    for doc in confidential_docs:
        print(f"    - {doc.metadata.get('source_file')} (dept: {doc.metadata.get('department')}, level: {doc.metadata.get('access_level')})")
    
    print("\nFiltering by access_level='public' (user can only see public):")
    public_docs = load_documents_by_access_level('public')
    print(f"  Found {len(public_docs)} documents")
    for doc in public_docs:
        print(f"    - {doc.metadata.get('source_file')} (dept: {doc.metadata.get('department')}, level: {doc.metadata.get('access_level')})")
    
    print("\nTest queries:")
    queries = get_test_queries()
    for scenario, query in queries.items():
        print(f"  {scenario}: '{query}'")
